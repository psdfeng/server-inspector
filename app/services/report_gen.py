"""
报告生成服务：支持 Word、HTML、PDF 格式
"""
from datetime import datetime
import os
import json


def generate_html_report(servers_data: list, report_date: str, config=None) -> str:
    """生成 HTML 格式巡检报告"""
    status_map = {'normal': ('正常', '#28a745', '✅'), 'warning': ('警告', '#ffc107', '⚠️'),
                  'critical': ('严重', '#dc3545', '🔴'), 'offline': ('离线', '#6c757d', '⚫'),
                  'unknown': ('未知', '#6c757d', '❓')}
    rows = ''
    for s in servers_data:
        status, color, icon = status_map.get(s.get('status', 'unknown'), ('未知', '#6c757d', '❓'))
        rows += f"""
        <tr>
            <td>{s.get('name','')}</td>
            <td>{s.get('ip','')}</td>
            <td>{s.get('os_label','')}</td>
            <td>{s.get('group','')}</td>
            <td style="color:{color};font-weight:bold">{icon} {status}</td>
            <td>{s.get('cpu_usage','N/A')}%</td>
            <td>{s.get('mem_usage','N/A')}%</td>
            <td>{s.get('max_disk_usage','N/A')}%</td>
            <td>{s.get('last_inspected','')}</td>
        </tr>"""
    
    total = len(servers_data)
    normal = sum(1 for s in servers_data if s.get('status') == 'normal')
    warning = sum(1 for s in servers_data if s.get('status') == 'warning')
    critical = sum(1 for s in servers_data if s.get('status') == 'critical')
    offline = sum(1 for s in servers_data if s.get('status') == 'offline')

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>服务器巡检报告 {report_date}</title>
<style>
  body{{font-family:'Microsoft YaHei',Arial,sans-serif;margin:20px;color:#333;background:#f8f9fa}}
  h1{{color:#1a3c5e;border-bottom:3px solid #1a3c5e;padding-bottom:10px}}
  h2{{color:#2c5f8a;margin-top:30px}}
  .summary{{display:flex;gap:15px;flex-wrap:wrap;margin:15px 0}}
  .card{{background:#fff;border-radius:8px;padding:15px 25px;text-align:center;box-shadow:0 2px 6px rgba(0,0,0,.1);min-width:120px}}
  .card .num{{font-size:2em;font-weight:bold}}
  .card .label{{color:#666;font-size:.9em}}
  .num-total{{color:#1a3c5e}} .num-normal{{color:#28a745}} .num-warning{{color:#ffc107}} .num-critical{{color:#dc3545}} .num-offline{{color:#6c757d}}
  table{{width:100%;border-collapse:collapse;background:#fff;border-radius:8px;overflow:hidden;box-shadow:0 2px 6px rgba(0,0,0,.1)}}
  th{{background:#1a3c5e;color:#fff;padding:10px 12px;text-align:left}}
  td{{padding:9px 12px;border-bottom:1px solid #eee}}
  tr:hover td{{background:#f0f7ff}}
  .footer{{margin-top:30px;text-align:center;color:#999;font-size:.85em}}
</style>
</head>
<body>
<h1>🏥 医院IT服务器巡检报告</h1>
<p>报告日期：<strong>{report_date}</strong> &nbsp;&nbsp; 生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>

<h2>📊 巡检概览</h2>
<div class="summary">
  <div class="card"><div class="num num-total">{total}</div><div class="label">服务器总数</div></div>
  <div class="card"><div class="num num-normal">{normal}</div><div class="label">正常</div></div>
  <div class="card"><div class="num num-warning">{warning}</div><div class="label">警告</div></div>
  <div class="card"><div class="num num-critical">{critical}</div><div class="label">严重</div></div>
  <div class="card"><div class="num num-offline">{offline}</div><div class="label">离线/未知</div></div>
</div>

<h2>📋 服务器巡检明细</h2>
<table>
  <thead><tr>
    <th>服务器名称</th><th>IP地址</th><th>系统类型</th><th>分组</th>
    <th>状态</th><th>CPU使用率</th><th>内存使用率</th><th>最大磁盘使用率</th><th>巡检时间</th>
  </tr></thead>
  <tbody>{rows}</tbody>
</table>
<div class="footer">本报告由医院IT服务器自动巡检系统生成</div>
</body></html>"""
    return html


def generate_word_report(servers_data: list, report_date: str) -> bytes:
    """生成 Word 格式巡检报告"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from io import BytesIO

    doc = Document()
    
    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # 标题
    title = doc.add_heading('医院 IT 服务器巡检报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'报告日期：{report_date}　　生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    # 概览
    doc.add_heading('巡检概览', level=1)
    total = len(servers_data)
    normal = sum(1 for s in servers_data if s.get('status') == 'normal')
    warning = sum(1 for s in servers_data if s.get('status') == 'warning')
    critical = sum(1 for s in servers_data if s.get('status') == 'critical')
    offline = sum(1 for s in servers_data if s.get('status') in ('offline', 'unknown'))
    
    summary_table = doc.add_table(rows=2, cols=5)
    summary_table.style = 'Table Grid'
    headers = ['服务器总数', '正常', '警告', '严重', '离线/未知']
    values = [str(total), str(normal), str(warning), str(critical), str(offline)]
    colors = [RGBColor(26, 60, 94), RGBColor(40, 167, 69), RGBColor(255, 193, 7), RGBColor(220, 53, 69), RGBColor(108, 117, 125)]
    for i, (h, v, c) in enumerate(zip(headers, values, colors)):
        cell = summary_table.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        val_cell = summary_table.cell(1, i)
        val_cell.text = v
        run = val_cell.paragraphs[0].runs[0]
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = c
    doc.add_paragraph('')

    # 明细表格
    doc.add_heading('服务器巡检明细', level=1)
    headers = ['名称', 'IP地址', '系统', '分组', '状态', 'CPU%', '内存%', '磁盘%', '巡检时间']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell._tc.get_or_add_tcPr().append(_make_shd_element('1A3C5E'))

    status_colors = {'normal': RGBColor(40, 167, 69), 'warning': RGBColor(255, 140, 0),
                     'critical': RGBColor(220, 53, 69), 'offline': RGBColor(108, 117, 125), 'unknown': RGBColor(108, 117, 125)}
    status_labels = {'normal': '正常', 'warning': '警告', 'critical': '严重', 'offline': '离线', 'unknown': '未知'}

    for s in servers_data:
        status = s.get('status', 'unknown')
        row = table.add_row()
        values = [
            s.get('name', ''), s.get('ip', ''), s.get('os_label', ''), s.get('group', ''),
            status_labels.get(status, '未知'),
            f"{s.get('cpu_usage', 'N/A')}%", f"{s.get('mem_usage', 'N/A')}%",
            f"{s.get('max_disk_usage', 'N/A')}%", s.get('last_inspected', '')
        ]
        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.text = str(val)
            if i == 4:  # 状态列着色
                cell.paragraphs[0].runs[0].font.color.rgb = status_colors.get(status, RGBColor(0, 0, 0))
                cell.paragraphs[0].runs[0].bold = True

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _make_shd_element(color_hex):
    """创建 Word 表格单元格背景色 XML 元素（w:shd）"""
    from docx.oxml.ns import qn
    from lxml import etree
    # 使用 qn() 生成正确的 Clark notation 标签名，如 {http://...}shd
    shd = etree.Element(qn('w:shd'))
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    return shd


def generate_pdf_report(servers_data: list, report_date: str, fonts_dir: str) -> bytes:
    """生成 PDF 格式巡检报告（ReportLab + 中文字体）"""
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from io import BytesIO

    # 注册中文字体：
    # 1) 优先使用 ReportLab 内置 CID 中文字体（最稳，不依赖本地 TTF/OTF 兼容）
    # 2) 再回退到本地字体文件
    font_name = 'Helvetica'
    try:
        pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
        font_name = 'STSong-Light'
    except Exception:
        pass

    if font_name == 'Helvetica':
        # 回退：注册本地中文字体（按候选列表自动匹配）
        # 注意：TTFont 对部分 OTF(CFF) 不兼容，因此保留 CID 作为主方案
        font_candidates = [
            ('SimHei', 'SimHei.ttf'),
            ('NotoSansSC', 'NotoSansSC-Regular.ttf'),
            ('SourceHanSansSC', 'SourceHanSansSC-Regular.ttf'),
            ('NotoSansSC', 'NotoSansSC-Regular.otf'),
            ('SourceHanSansSC', 'SourceHanSansSC-Regular.otf'),
            ('SourceHanSansCN', 'SourceHanSansCN-Regular.otf'),
        ]
        for alias, filename in font_candidates:
            path = os.path.join(fonts_dir, filename)
            if not os.path.exists(path):
                continue
            try:
                pdfmetrics.registerFont(TTFont(alias, path))
                font_name = alias
                break
            except Exception:
                continue

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4),
                            leftMargin=1.5*cm, rightMargin=1.5*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('ChineseTitle', fontName=font_name, fontSize=18, spaceAfter=10,
                                  alignment=1, textColor=colors.HexColor('#1A3C5E'))
    normal_style = ParagraphStyle('ChineseNormal', fontName=font_name, fontSize=10, spaceAfter=6)
    
    story = []
    story.append(Paragraph('医院 IT 服务器巡检报告', title_style))
    story.append(Paragraph(f'报告日期：{report_date}　　生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', normal_style))
    story.append(Spacer(1, 0.5*cm))

    # 概览
    total = len(servers_data)
    normal = sum(1 for s in servers_data if s.get('status') == 'normal')
    warning = sum(1 for s in servers_data if s.get('status') == 'warning')
    critical = sum(1 for s in servers_data if s.get('status') == 'critical')
    offline = total - normal - warning - critical

    story.append(Paragraph('巡检概览', ParagraphStyle('H1', fontName=font_name, fontSize=14, spaceAfter=8, textColor=colors.HexColor('#1A3C5E'))))
    summary_data = [
        ['服务器总数', '正常', '警告', '严重', '离线/未知'],
        [str(total), str(normal), str(warning), str(critical), str(offline)]
    ]
    summary_table = Table(summary_data, colWidths=[4*cm]*5)
    summary_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), font_name),
        ('FONTSIZE', (0, 0), (-1, 0), 11), ('FONTSIZE', (0, 1), (-1, 1), 16),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1A3C5E')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('TEXTCOLOR', (0, 1), (0, 1), colors.HexColor('#1A3C5E')),
        ('TEXTCOLOR', (1, 1), (1, 1), colors.HexColor('#28a745')),
        ('TEXTCOLOR', (2, 1), (2, 1), colors.HexColor('#ffc107')),
        ('TEXTCOLOR', (3, 1), (3, 1), colors.HexColor('#dc3545')),
        ('TEXTCOLOR', (4, 1), (4, 1), colors.HexColor('#6c757d')),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTSIZE', (0, 1), (-1, 1), 18), ('FONTNAME', (0, 1), (-1, 1), font_name),
        ('BOLD', (0, 1), (-1, 1), True),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('ROWBACKGROUNDS', (0, 1), (-1, 1), [colors.HexColor('#f8f9fa')]),
        ('TOPPADDING', (0, 0), (-1, -1), 8), ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 0.5*cm))

    # 明细
    story.append(Paragraph('服务器巡检明细', ParagraphStyle('H1', fontName=font_name, fontSize=14, spaceAfter=8, textColor=colors.HexColor('#1A3C5E'))))
    table_data = [['名称', 'IP地址', '系统', '分组', '状态', 'CPU%', '内存%', '磁盘%', '巡检时间']]
    status_labels = {'normal': '正常', 'warning': '警告', 'critical': '严重', 'offline': '离线', 'unknown': '未知'}
    for s in servers_data:
        table_data.append([
            s.get('name', ''), s.get('ip', ''), s.get('os_label', ''), s.get('group', ''),
            status_labels.get(s.get('status', 'unknown'), '未知'),
            f"{s.get('cpu_usage', 'N/A')}%", f"{s.get('mem_usage', 'N/A')}%",
            f"{s.get('max_disk_usage', 'N/A')}%", s.get('last_inspected', ''),
        ])
    
    col_widths = [4*cm, 3.5*cm, 2.5*cm, 3*cm, 2*cm, 2*cm, 2*cm, 2*cm, 3.5*cm]
    detail_table = Table(table_data, colWidths=col_widths, repeatRows=1)
    style_cmds = [
        ('FONTNAME', (0, 0), (-1, -1), font_name),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1A3C5E')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.3, colors.grey),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('TOPPADDING', (0, 0), (-1, -1), 5), ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
    ]
    status_bg = {'normal': colors.HexColor('#d4edda'), 'warning': colors.HexColor('#fff3cd'),
                 'critical': colors.HexColor('#f8d7da'), 'offline': colors.HexColor('#e2e3e5'), 'unknown': colors.HexColor('#e2e3e5')}
    for i, s in enumerate(servers_data, 1):
        bg = status_bg.get(s.get('status', 'unknown'), colors.white)
        style_cmds.append(('BACKGROUND', (4, i), (4, i), bg))
    detail_table.setStyle(TableStyle(style_cmds))
    story.append(detail_table)
    
    doc.build(story)
    buf.seek(0)
    return buf.read()
